import os
import pytest
import docx
from docx.document import Document as DocxDocument

from enhanced_md import EnhancedMD

# ----- PYTEST FIXTURES -----

# # ----- __init__ fixtures -----

def set_test_docx_styles(docx_doc: DocxDocument):
	# Create heading styles (with 3 clearly defined hierarchy levels and 1 undefined)
	docx_doc.styles.add_style(name="test_h0", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	docx_doc.styles.add_style(name="test_h1", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	docx_doc.styles.add_style(name="test_h2", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	docx_doc.styles.add_style(name="test_h3", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

	# Create paragraph styles (with 3 clearly defined hierarchy levels and 1 undefined)
	docx_doc.styles.add_style(name="test_p0", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	docx_doc.styles.add_style(name="test_p1", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	docx_doc.styles.add_style(name="test_p2", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
	docx_doc.styles.add_style(name="test_p3", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

	# Create an undefined style
	docx_doc.styles.add_style(name="test_undefined", style_type=docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)


@pytest.fixture
def create_empty_test_docx_document():
	# Set up: Create test docx document
	docx_doc = docx.Document()
	set_test_docx_styles(docx_doc=docx_doc)

	docx_file_path = "test.docx"
	docx_doc.save(docx_file_path)

	yield docx_doc, docx_file_path

	# Tear down: Delete test docx_document
	os.remove(docx_file_path)


@pytest.fixture
def create_test_styles_dict():
	styles = {
		"heading": {
			0: ["test_h0", "test_undefined"],
			1: ["test_h1"],
			2: ["test_h2"],
			3: ["test_h3"]
		},
		"paragraph": {
			0: ["test_p0", "test_undefined"],
			1: ["test_p1"],
			2: ["test_p2"],
			3: ["test_p3"]
		}
	}

	yield styles

# # ----- build_doc_graph fixtures -----

@pytest.fixture(params=[0, 1, 2, 3])
def fill_test_docx_document_with_starting_heading(request, create_empty_test_docx_document):
	# Set up:
	docx_doc, docx_file_path = create_empty_test_docx_document
	docx_doc.add_paragraph(text="start H", style=f"test_h{request.param}")
	docx_doc.add_paragraph(text="H", style="test_h1")  # Followed by a heading to check correct doc graph root append
	docx_doc.save(docx_file_path)

	yield docx_file_path

	# Tear down:
	docx_doc = docx.Document()
	docx_doc.save(docx_file_path)


@pytest.fixture(params=[0, 1, 2, 3])
def fill_test_docx_document_with_starting_paragraph(request, create_empty_test_docx_document):
	# Set up:
	docx_doc, docx_file_path = create_empty_test_docx_document
	docx_doc.add_paragraph(text="start P", style=f"test_p{request.param}")
	docx_doc.add_paragraph(text="H", style="test_h1")  # Followed by a heading to check correct doc graph root append
	docx_doc.save(docx_file_path)

	yield docx_file_path

	# Tear down:
	docx_doc = docx.Document()
	docx_doc.save(docx_file_path)


@pytest.fixture
def fill_test_docx_document_with_starting_undefined(create_empty_test_docx_document):
	docx_doc, docx_file_path = create_empty_test_docx_document
	docx_doc.add_paragraph(text="start U", style=f"test_undefined")
	docx_doc.add_paragraph(text="H", style="test_h1")  # Followed by a heading to check correct doc graph root append
	docx_doc.save(docx_file_path)

	yield docx_file_path

	# Tear down:
	docx_doc = docx.Document()
	docx_doc.save(docx_file_path)

# ----- UNIT TESTS -----

# # ----- build_doc_graph -----

def test_build_doc_graph_with_starting_heading(fill_test_docx_document_with_starting_heading, create_test_styles_dict):
	#
	docx_file_path = fill_test_docx_document_with_starting_heading
	styles = create_test_styles_dict

	#
	test_emd = EnhancedMD(docx_file_path=docx_file_path, styles=styles)
	test_emd()

	# Ensure correct doc graph roots structure
	assert len(test_emd.doc_graph) == 2
	assert test_emd.doc_graph[0].parent is None
	assert test_emd.doc_graph[1].parent is None

	assert test_emd.doc_graph[0].children == []  # Ensure second heading is not assigned as children of the first heading
	assert test_emd.doc_graph[1].children == []  # Ensure first heading is not assigned as children of the second heading

	# Ensure correct heading next and previous relations
	assert test_emd.doc_graph[0].next == test_emd.doc_graph[1]
	assert test_emd.doc_graph[1].previous == test_emd.doc_graph[0]


def test_build_doc_graph_with_starting_paragraph(fill_test_docx_document_with_starting_paragraph,
                                                 create_test_styles_dict):
	#
	docx_file_path = fill_test_docx_document_with_starting_paragraph
	styles = create_test_styles_dict

	#
	test_emd = EnhancedMD(docx_file_path=docx_file_path, styles=styles)
	test_emd()

	# Ensure correct doc graph roots structure
	assert len(test_emd.doc_graph) == 2
	assert test_emd.doc_graph[0].parent is None
	assert test_emd.doc_graph[1].parent is None

	assert test_emd.doc_graph[0].children == []  # Ensure heading is not assigned as children of paragraph
	assert test_emd.doc_graph[1].children == []  # Ensure paragraph is not assigned as children of heading

	# Ensure correct paragraph and heading next and previous relations
	assert test_emd.doc_graph[0].next == test_emd.doc_graph[1]
	assert test_emd.doc_graph[1].previous == test_emd.doc_graph[0]


def test_build_doc_graph_with_starting_undefined(fill_test_docx_document_with_starting_undefined,
                                                 create_test_styles_dict):
	#
	docx_file_path = fill_test_docx_document_with_starting_undefined
	styles = create_test_styles_dict

	#
	test_emd = EnhancedMD(docx_file_path=docx_file_path, styles=styles)
	test_emd()

	# Ensure correct doc graph roots structure
	assert len(test_emd.doc_graph) == 2
	assert test_emd.doc_graph[0].parent is None
	assert test_emd.doc_graph[1].parent is None

	assert test_emd.doc_graph[0].children == []  # Ensure heading is not assigned as children of undefined
	assert test_emd.doc_graph[1].children == []  # Ensure undefined is not assigned as children of the heading

	# Ensure correct undefined and heading next and previous relations
	assert test_emd.doc_graph[0].next == test_emd.doc_graph[1]
	assert test_emd.doc_graph[1].previous == test_emd.doc_graph[0]
